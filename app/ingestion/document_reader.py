# app/ingestion/document_reader.py

import os
import logging
from typing import Optional, Dict, Any, List, Union
import mimetypes

logger = logging.getLogger(__name__)

class DocumentReader:
    """
    Clase para leer documentos de diferentes formatos y extraer su texto.
    Soporta múltiples formatos como PDF, TXT, DOCX, etc.
    """
    
    def __init__(self):
        # Registrar tipos MIME para diferentes formatos
        mimetypes.init()
        # Asegurar que los formatos comunes estén registrados
        if not mimetypes.guess_type('test.pdf')[0]:
            mimetypes.add_type('application/pdf', '.pdf')
        if not mimetypes.guess_type('test.docx')[0]:
            mimetypes.add_type('application/vnd.openxmlformats-officedocument.wordprocessingml.document', '.docx')
    
    def read_document(self, file_path: str) -> Dict[str, Any]:
        """
        Lee un documento y extrae su texto basado en su formato.
        
        Args:
            file_path: Ruta al archivo a leer
            
        Returns:
            Diccionario con el contenido y metadatos del documento
            
        Raises:
            ValueError: Si el formato no es soportado o el archivo no existe
            ImportError: Si faltan dependencias para leer el formato
        """
        if not os.path.exists(file_path):
            raise ValueError(f"El archivo no existe: {file_path}")
        
        # Determinar el tipo de archivo basado en la extensión
        mime_type, _ = mimetypes.guess_type(file_path)
        
        if not mime_type:
            # Intentar inferir por extensión si mime_type falla
            _, extension = os.path.splitext(file_path)
            extension = extension.lower()
            
            if extension == '.pdf':
                return self._read_pdf(file_path)
            elif extension in ['.txt', '.md']:
                return self._read_text(file_path)
            elif extension in ['.docx', '.doc']:
                return self._read_word(file_path)
            else:
                raise ValueError(f"Formato de archivo no soportado: {extension}")
        
        # Procesar según el tipo MIME
        if mime_type == 'application/pdf':
            return self._read_pdf(file_path)
        elif mime_type.startswith('text/'):
            return self._read_text(file_path)
        elif mime_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                           'application/msword']:
            return self._read_word(file_path)
        else:
            raise ValueError(f"Tipo MIME no soportado: {mime_type}")
    
    def _read_pdf(self, file_path: str) -> Dict[str, Any]:
        """
        Lee un archivo PDF y extrae su texto usando PyPDF2.
        
        Args:
            file_path: Ruta al archivo PDF
            
        Returns:
            Dict con contenido y metadatos
        """
        try:
            from PyPDF2 import PdfReader
        except ImportError:
            raise ImportError("PyPDF2 es necesario para extraer texto de archivos PDF. "
                             "Instálalo con: pip install PyPDF2")
        
        try:
            reader = PdfReader(file_path)
            
            # Extraer texto
            text = ""
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            
            # Extraer metadatos disponibles
            metadata = {}
            if reader.metadata:
                for key in reader.metadata:
                    if reader.metadata[key]:
                        # Convertir ByteString a str si es necesario
                        value = reader.metadata[key]
                        if isinstance(value, bytes):
                            try:
                                value = value.decode('utf-8')
                            except UnicodeDecodeError:
                                value = str(value)
                        metadata[key] = value
            
            return {
                'content': text,
                'metadata': metadata,
                'pages': len(reader.pages),
                'format': 'pdf'
            }
        except Exception as e:
            logger.error(f"Error leyendo PDF {file_path}: {e}")
            raise
    
    def _read_text(self, file_path: str) -> Dict[str, Any]:
        """
        Lee un archivo de texto plano.
        
        Args:
            file_path: Ruta al archivo de texto
            
        Returns:
            Dict con contenido y metadatos básicos
        """
        try:
            encodings = ['utf-8', 'latin-1', 'cp1252']
            text = None
            
            # Intentar diferentes codificaciones
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text = file.read()
                    break
                except UnicodeDecodeError:
                    continue
            
            if text is None:
                raise ValueError(f"No se pudo decodificar el archivo {file_path} con ninguna codificación conocida")
            
            return {
                'content': text,
                'metadata': {
                    'filename': os.path.basename(file_path),
                    'size': os.path.getsize(file_path)
                },
                'format': 'text'
            }
        except Exception as e:
            logger.error(f"Error leyendo archivo de texto {file_path}: {e}")
            raise
    
    def _read_word(self, file_path: str) -> Dict[str, Any]:
        """
        Lee un archivo Word (.docx) y extrae su texto usando python-docx.
        
        Args:
            file_path: Ruta al archivo Word
            
        Returns:
            Dict con contenido y metadatos
        """
        try:
            import docx
        except ImportError:
            raise ImportError("python-docx es necesario para extraer texto de archivos Word. "
                             "Instálalo con: pip install python-docx")
        
        try:
            doc = docx.Document(file_path)
            
            # Extraer texto de párrafos
            text = "\n".join([para.text for para in doc.paragraphs if para.text])
            
            # Extraer texto de tablas
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells if cell.text])
                    if row_text:
                        text += "\n" + row_text
            
            # Extraer metadatos disponibles
            metadata = {
                'filename': os.path.basename(file_path),
                'size': os.path.getsize(file_path)
            }
            
            if hasattr(doc, 'core_properties'):
                props = doc.core_properties
                for attr in ['author', 'category', 'comments', 'content_status', 
                            'created', 'identifier', 'keywords', 'language', 
                            'last_modified_by', 'last_printed', 'modified', 
                            'revision', 'subject', 'title', 'version']:
                    if hasattr(props, attr):
                        value = getattr(props, attr)
                        if value:
                            metadata[attr] = str(value)
            
            return {
                'content': text,
                'metadata': metadata,
                'format': 'docx'
            }
        except Exception as e:
            logger.error(f"Error leyendo Word {file_path}: {e}")
            raise
    
    def read_multiple(self, directory: str, extensions: List[str] = None, recursive: bool = True) -> List[Dict[str, Any]]:
        """
        Lee múltiples documentos de un directorio, opcionalmente filtrando por extensiones.
        
        Args:
            directory: Directorio a leer
            extensions: Lista de extensiones a incluir (ej: ['.pdf', '.txt'])
            recursive: Si se deben buscar archivos en subdirectorios
            
        Returns:
            Lista de documentos procesados
        """
        if not os.path.isdir(directory):
            raise ValueError(f"El directorio no existe: {directory}")
        
        results = []
        
        # Normalizar extensiones
        if extensions:
            extensions = [ext.lower() if ext.startswith('.') else f'.{ext.lower()}' for ext in extensions]
        
        for root, _, files in os.walk(directory):
            if not recursive and root != directory:
                continue
                
            for file in files:
                file_path = os.path.join(root, file)
                _, ext = os.path.splitext(file_path)
                
                if extensions and ext.lower() not in extensions:
                    continue
                
                try:
                    document = self.read_document(file_path)
                    # Añadir ruta relativa para referencia
                    document['file_path'] = os.path.relpath(file_path, directory)
                    document['title'] = os.path.basename(file_path)
                    results.append(document)
                except Exception as e:
                    logger.warning(f"Error procesando {file_path}: {e}")
                    continue
        
        return results